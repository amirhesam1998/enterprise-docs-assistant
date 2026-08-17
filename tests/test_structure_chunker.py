from __future__ import annotations

from datetime import datetime, timezone

import docx
import openpyxl
import pytest

from eda.extractors.base import ExtractedDocument, ExtractionOptions, ExtractorRoute
from eda.extractors.docx import DocxExtractor
from eda.extractors.xlsx import XlsxExtractor
from eda.ingestion_schema import (
    DocumentManifest,
    ExtractionRoute as PageRoute,
    OCRDecision,
    ProcessingStatus,
)
from eda.structure_chunker import structure_chunks


@pytest.fixture
def docx_document(tmp_path):
    document = docx.Document()
    document.add_heading("Quarterly Report", level=1)
    document.add_paragraph(
        "Revenue increased by twelve percent this quarter across every region and product line."
    )
    document.add_paragraph("Operating costs fell after automation and vendor renegotiation.")
    table = document.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Category"
    table.rows[0].cells[1].text = "Amount"
    table.rows[1].cells[0].text = "Rent"
    table.rows[1].cells[1].text = "1200"
    table.rows[2].cells[0].text = "Food"
    table.rows[2].cells[1].text = "300"
    path = tmp_path / "report.docx"
    document.save(path)
    options = ExtractionOptions(
        tenant_id="kb", logical_document_key="report.docx",
        source_name="report.docx", acl_groups=("billing",),
    )
    return DocxExtractor().extract(path, options=options)


def test_chunks_preserve_tenant_acl_and_provenance(docx_document):
    chunks = structure_chunks(docx_document, acl_groups=["billing", "billing"])
    assert chunks
    for chunk in chunks:
        assert chunk.tenant_id == "kb"           # from the authenticated uploader, via manifest
        assert chunk.acl_groups == ["billing"]   # authoritative + de-duplicated
        assert chunk.source == "report.docx"
        loc = chunk.location
        assert loc["document_id"] and loc["revision_id"] and loc["page_id"]
        assert loc["extractor_route"] == "docx"
        assert "block_ids" in loc


def test_table_chunk_keeps_row_and_column_context(docx_document):
    chunks = structure_chunks(docx_document, acl_groups=["billing"])
    table_chunks = [c for c in chunks if c.location["kind"] == "table"]
    assert len(table_chunks) == 1
    text = table_chunks[0].text
    assert "Category: Rent" in text and "Amount: 1200" in text


def test_chunk_ids_are_deterministic(tmp_path):
    def extract():
        document = docx.Document()
        document.add_paragraph("A stable paragraph used to check deterministic identifiers.")
        path = tmp_path / "stable.docx"
        document.save(path)
        options = ExtractionOptions(
            tenant_id="kb", logical_document_key="stable.docx",
            source_name="stable.docx", acl_groups=("billing",),
        )
        return structure_chunks(DocxExtractor().extract(path, options=options), acl_groups=["billing"])

    assert [c.chunk_id for c in extract()] == [c.chunk_id for c in extract()]


def test_missing_acl_is_rejected(docx_document):
    with pytest.raises(ValueError, match="ACL"):
        structure_chunks(docx_document, acl_groups=[])


def test_xlsx_requires_authoritative_acl(tmp_path):
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.append(["Category", "Amount"])
    sheet.append(["Rent", 1200])
    path = tmp_path / "budget.xlsx"
    workbook.save(path)

    with pytest.raises(ValueError, match="ACL"):
        XlsxExtractor().extract(
            path,
            options=ExtractionOptions(
                tenant_id="kb", logical_document_key="budget.xlsx",
                source_name="budget.xlsx", acl_groups=(),
            ),
        )
    document = XlsxExtractor().extract(
        path,
        options=ExtractionOptions(
            tenant_id="kb", logical_document_key="budget.xlsx",
            source_name="budget.xlsx", acl_groups=("billing",),
        ),
    )
    chunks = structure_chunks(document, acl_groups=["billing"])
    assert any("Category: Rent" in c.text for c in chunks)


# --- low-quality pages are never silently embedded ------------------------

def _manifest(identity) -> DocumentManifest:
    return DocumentManifest(
        document_id=identity["document_id"],
        revision_id=identity["revision_id"],
        tenant_id="tenant-a",
        source_name="doc.pdf",
        source_type="pdf",
        source_sha256=identity["digest"],
        file_size_bytes=10,
        mime_type="application/pdf",
        total_pages=1,
        parser_name="fake",
        parser_version="1.0",
        created_at=datetime(2026, 1, 1, tzinfo=timezone.utc),
    )


def _doc(identity, page) -> ExtractedDocument:
    return ExtractedDocument.build(
        manifest=_manifest(identity), pages=[page], extractor_route=ExtractorRoute.NATIVE_PDF,
    )


def test_needs_review_page_not_chunked_unless_allowed(identity, page_factory):
    page = page_factory(
        native_text="This page needs a human review before it can be trusted for retrieval use.",
        processing_status=ProcessingStatus.NEEDS_REVIEW,
        ocr_decision=OCRDecision(
            selected_route=PageRoute.NATIVE, quality_gate_passed=False, candidate_count=0
        ),
    )
    document = _doc(identity, page)
    assert structure_chunks(document, acl_groups=["billing"]) == []
    assert structure_chunks(document, acl_groups=["billing"], allow_needs_review=True)


def test_failed_page_is_never_chunked(identity, page_factory):
    from eda.ingestion_schema import PageType

    page = page_factory(
        native_text="",
        processing_status=ProcessingStatus.FAILED,
        page_type=PageType.EMPTY,
        ocr_decision=OCRDecision(
            selected_route=PageRoute.REJECTED, quality_gate_passed=False, candidate_count=0
        ),
    )
    document = _doc(identity, page)
    assert structure_chunks(document, acl_groups=["billing"], allow_needs_review=True) == []
